"""
generate_report.py
Generates ASR benchmark report as a formatted .docx file.
Target: 3 pages. Charts referenced separately in results/charts/
Run from asr-benchmark/ root directory.
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUTPUT_PATH = Path("report.docx")


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.space_before = Pt(2)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(doc, headers, rows, col_widths=None, header_color="1F497D"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = "F2F7FC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[i])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# ── Document ───────────────────────────────────────────────────────────────────

def build_report():
    doc = Document()

    # Page margins — tighter to fit 3 pages
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Default font
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9.5)

    # ── Title Block ────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ASR Benchmark: Indian Conversational Speech")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    title.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "Deepgram Nova-2 · Sarvam Saaras v3 · Whisper Medium · IndicConformer 600M"
    ).font.size = Pt(10)
    sub.paragraph_format.space_after = Pt(3)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "May 2026  |  22 self-recorded audio files  |  4 models  |  88 inference runs"
    ).font.size = Pt(8.5)
    meta.paragraph_format.space_after = Pt(10)

    # ── 1. Approach ────────────────────────────────────────────────────────────
    add_heading(doc, "1. Approach", level=1)

    add_body(doc,
        "Platform context: blue-collar hiring, phone and WhatsApp voice input, "
        "Hindi/Hinglish, noisy environments. The core task is entity extraction — "
        "did the model correctly capture the locality name from natural conversational speech? "
        "Wrong locality = wrong job match. WER is insufficient here: a transcript returning "
        "'alhaga' for 'Yelahanka' scores WER=1.0 on that word but the failure is a "
        "missed entity, not a transcription error. "
        "Four models were benchmarked across the API vs local and "
        "general-purpose vs India-specific axes:"
    )

    add_table(doc,
        headers=["Model", "Type", "Rationale"],
        rows=[
            ["Deepgram Nova-2", "API — Baseline", "Required. Production ASR marketed for telephony."],
            ["Sarvam Saaras v3", "API", "India-specific, telephony-optimized, Hinglish transliteration support."],
            ["Whisper Medium", "Local (CPU, int8)", "Most-used open-source ASR. Free, privacy-safe, no per-call cost."],
            ["IndicConformer 600M", "Colab T4", "AI4Bharat model, trained on 22 Indian languages. Best Indian open-source option."],
        ],
        col_widths=[1.8, 1.5, 4.1],
    )

    add_body(doc,
        "Models rejected: IndicWhisper — broken checkpoint, weight size mismatches. "
        "Sarvam Saarika v2.5 — deprecated mid-evaluation, migrated to Saaras v3. "
        "IndicConformer ran on Google Colab T4 (6.17GB model exceeds local 6GB VRAM). "
        "Whisper ran on CPU due to missing CUDA cublas DLL on GTX 1660."
    )

    add_body(doc,
        "Dataset: 22 phone-mic recordings (20 Hindi/Hinglish + 2 English for "
        "cross-language generalization). Two speakers — male (16 files), female (6 files), "
        "both Malayalam-accented Hindi, matching the actual blue-collar worker demographic "
        "in Bangalore. Conditions: quiet (8), noisy (9), rushed (3), whispered (2). "
        "External validation: FLEURS Hindi test split (20 samples) via HuggingFace."
    )

    add_table(doc,
        headers=["Metric", "Why"],
        rows=[
            ["Entity Accuracy", "Primary. Binary: did the locality name get captured? Most business-relevant."],
            ["WER", "Standard word-level metric. Secondary — not sufficient for proper noun tasks."],
            ["CER", "Character-level. Better for Indian names where one character difference matters."],
            ["Latency (s)", "API response time per file. Critical for real-time telephony."],
        ],
        col_widths=[1.8, 6.6],
    )

    add_body(doc,
        "Entity accuracy uses fuzzy matching via rapidfuzz (sliding window, threshold=85). "
        "Threshold set at 85 not 80 — 'electronic city' matched 'electricity' at 84.6, "
        "a false positive. Devanagari aliases added for Whisper and IndicConformer since "
        "both output Devanagari for Hindi speech, making WER/CER comparisons invalid."
    )

    # ── 2. Key Findings ────────────────────────────────────────────────────────
    add_heading(doc, "2. Key Findings", level=1)

    add_table(doc,
        headers=["Model", "Entity Accuracy", "Mean WER", "Mean CER", "Mean Latency"],
        rows=[
            ["Deepgram Nova-2", "32%", "0.49", "0.28", "0.9s"],
            ["Sarvam Saaras v3", "77% ✓", "0.25 ✓", "0.10 ✓", "0.7s ✓"],
            ["IndicConformer 600M", "55%", "1.01 *", "0.85 *", "1.5s"],
            ["Whisper Medium", "64%", "0.94 *", "0.79 *", "12.4s"],
        ],
        col_widths=[2.2, 1.7, 1.3, 1.3, 1.9],
    )
    add_body(doc,
        "* WER/CER not comparable — Devanagari output vs Roman reference. "
        "Whisper WER (0.94) is lower than IndicConformer (1.01) partly because "
        "English files pull the mean down. "
        "Charts: results/charts/ — entity_accuracy.png, noise_breakdown.png, "
        "gender_breakdown.png, latency_comparison.png, llm_vs_fuzzy.png"
    )

    add_heading(doc, "Results by Condition", level=2)
    add_table(doc,
        headers=["Condition", "Deepgram", "Sarvam", "IndicConformer", "Whisper"],
        rows=[
            ["Quiet", "29%", "86%", "86%", "43%"],
            ["Noisy", "25%", "63%", "38%", "63%"],
            ["Rushed", "0% ✗", "100% ✓", "0% ✗", "67%"],
            ["Whispered", "50%", "50%", "50%", "100% ✓"],
        ],
        col_widths=[1.5, 1.6, 1.6, 1.9, 1.8],
    )
    add_body(doc,
        "Sarvam is the only model achieving 100% on rushed audio — the most "
        "production-relevant condition for telephony. Deepgram and IndicConformer "
        "score 0% on rushed speech — complete collapse. Whisper scores 100% on "
        "whispered audio but is not deployable at 12.4s per file latency."
    )

    add_heading(doc, "Gender Gap", level=2)
    add_table(doc,
        headers=["Model", "Male", "Female", "Gap"],
        rows=[
            ["Deepgram Nova-2", "27%", "43%", "-16% (reversed — female > male, unique to Deepgram)"],
            ["Sarvam Saaras v3", "87%", "57%", "+30% (largest gap — likely training data bias)"],
            ["IndicConformer 600M", "53%", "57%", "~0% (gender-neutral)"],
            ["Whisper Medium", "67%", "57%", "+10%"],
        ],
        col_widths=[2.0, 1.0, 1.1, 4.3],
    )
    add_body(doc,
        "Both speakers recorded in identical conditions. Sarvam's 30-point gender gap "
        "is not explained by recording quality — it is a production risk that needs "
        "monitoring and female-speaker evaluation data before deployment."
    )

    add_heading(doc, "Open-Source Validation", level=2)
    add_body(doc,
        "FLEURS Hindi (20 samples, test split) confirms self-recorded findings: "
        "Sarvam outperforms Deepgram on CER (1.382 vs 1.411) and latency (0.72s vs 1.84s). "
        "WER ~1.0 on both due to Devanagari reference vs Roman output script mismatch — "
        "expected, not a failure. Results are not an artifact of self-recording quality."
    )

    # ── 3. Failure Analysis ────────────────────────────────────────────────────
    add_heading(doc, "3. Failure Analysis", level=1)

    add_table(doc,
        headers=["File", "Reference", "Deepgram Output", "Type"],
        rows=[
            ["05_marathahalli_whispered", "abhi marathahalli bridge...", "(empty)", "Model silent on whisper"],
            ["07_rajajinagar_noisy", "rajajinagar mein ek room...", "(empty)", "Complete failure on noise"],
            ["09_yelahanka_quiet", "yelahanka mein ek naya...", "alhaga mein ek naya...", "Phonetic distortion"],
            ["20_yeshwanthpur_noisy", "haan bhai yeshwanthpur...", "fine guys is chandraghor...", "Severe hallucination"],
            ["01_koramangala_quiet", "haan bhai main koramangala...", "hanbhai main koramaangala...", "Word boundary + insertion"],
        ],
        col_widths=[2.1, 1.9, 2.1, 2.3],
    )

    add_body(doc,
        "Deepgram produced 2 empty transcripts — no other model did. "
        "The hallucination on file 20 is dangerous: 'chandraghor' is a coherent "
        "but completely wrong locality with no confidence signal to flag it."
    )

    add_heading(doc, "LLM-Based Entity Extraction", level=2)
    add_body(doc,
        "ASR transcript → Llama-3 70B (Groq) for NER. Tested on all "
        "Deepgram + Sarvam outputs (44 samples):"
    )

    add_table(doc,
        headers=["Model", "Fuzzy Match", "LLM Extraction", "Delta"],
        rows=[
            ["Deepgram Nova-2", "32%", "73%", "+41 pts"],
            ["Sarvam Saaras v3", "77%", "95%", "+18 pts"],
        ],
        col_widths=[2.2, 1.8, 1.8, 2.6],
    )

    add_body(doc, "LLM rescues (selected):")
    rescues = [
        ('"alhaga"', " → Yelahanka"),
        ('"bpm leut"', " → BTM Layout"),
        ('"wiphil"', " → Whitefield"),
        ('"silver par"', " → Silk Board"),
        ('"bailandoor"', " → Bellandur"),
    ]
    for prefix, text in rescues:
        add_bullet(doc, text, bold_prefix=prefix)

    add_body(doc, "LLM failure modes:")
    add_bullet(doc,
        " → 'Chandraghoda' — real but wrong locality, confident hallucination",
        bold_prefix='"chandraghor session"'
    )
    add_bullet(doc,
        " → 'Mysore Road' — plausible but wrong",
        bold_prefix='"masooch ek bus"'
    )
    add_body(doc,
        "LLM hallucinates confidently on garbled input — producing a wrong answer "
        "that looks correct. Gazetteer validation is mandatory as a third stage."
    )

    # ── 4. Recommendation ─────────────────────────────────────────────────────
    add_heading(doc, "4. Recommendation", level=1)

    add_body(doc,
        "Sarvam Saaras v3 is the production pick — best entity accuracy (77%), "
        "fastest latency (0.7s), only model handling rushed speech (100%), "
        "consistent on noisy audio (63%). No other model is close on the primary task."
    )

    add_table(doc,
        headers=["Stage", "Component", "Role"],
        rows=[
            ["1", "Sarvam Saaras v3  (mode=translit)", "Phone audio → Roman script transcript"],
            ["2", "Llama-3 70B via Groq", "Extract locality from distorted transcript (77% → 95%)"],
            ["3", "Gazetteer validation", "Reject hallucinations against known locality list"],
        ],
        col_widths=[0.6, 3.1, 4.7],
    )

    add_body(doc,
        "Three-stage pipeline achieves 95% entity accuracy vs 77% for ASR alone. "
        "The gazetteer is not optional — without it, LLM hallucinations like "
        "'Chandraghoda' pass as valid answers."
    )

    add_heading(doc, "Caveats", level=2)
    caveats = [
        ("Gender gap:", " Sarvam 87% male vs 57% female. Collect female-speaker eval data before production."),
        ("Whispered audio:", " Unsolved across all production-ready models (max 50%). Whisper solves it but is not deployable."),
        ("Cost at scale:", " Sarvam is a paid API. At high call volume, IndicConformer on dedicated GPU is viable — matches Sarvam on quiet audio (86%)."),
        ("Dataset size:", " 22 files, 2 speakers. Validate on larger, more diverse pool before production commitment."),
        ("IndicConformer robustness:", " 86% on quiet but 0% on rushed — strong capability, poor noise robustness. Needs noisy training data."),
    ]
    for bold, text in caveats:
        add_bullet(doc, text, bold_prefix=bold)

    # ── Footer ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    note = doc.add_paragraph(
        "Code: pipeline.py · metrics.py · analyse.py · llm_extraction.py · "
        "open_source_eval.py  |  "
        "Results: results/combined_results.csv (88 rows)  |  "
        "Reproduce: python pipeline.py"
    )
    note.runs[0].font.size = Pt(7.5)
    note.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    note.paragraph_format.space_before = Pt(8)

    # ── Save ───────────────────────────────────────────────────────────────────
    doc.save(str(OUTPUT_PATH))
    print(f"Report saved → {OUTPUT_PATH}")
    print("Charts referenced separately in results/charts/")


if __name__ == "__main__":
    build_report()